"""
DOCX File Handler
Translates Word documents while preserving:
- Paragraph formatting (font, size, bold, italic, etc.)
- Tables and their content
- Headers and footers
- Document structure (sections, page breaks)
"""

from docx import Document
from core.utils import is_translatable


class DocxHandler:
    """Handles translation of .docx files."""

    def translate(self, input_path, output_path, translator, progress_callback=None):
        """
        Translate a DOCX file.

        Strategy:
        - For each paragraph: combine all run texts, translate, 
          put result in first run, clear other runs.
          This preserves the paragraph-level formatting from the first run.
        - Tables, headers, and footers are handled similarly.
        """
        doc = Document(input_path)

        # Count total translatable elements for progress
        total_elements = self._count_elements(doc)
        processed = 0

        # 1. Translate body paragraphs
        for para in doc.paragraphs:
            if is_translatable(para.text):
                translated = translator.translate_text(para.text)
                self._set_paragraph_text(para, translated)
            processed += 1
            if progress_callback and processed % 5 == 0:
                progress_callback(
                    processed / total_elements,
                    f"Translating paragraphs... ({processed}/{total_elements})"
                )

        # 2. Translate tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if is_translatable(para.text):
                            translated = translator.translate_text(para.text)
                            self._set_paragraph_text(para, translated)
                        processed += 1

            if progress_callback:
                progress_callback(
                    min(processed / total_elements, 0.95),
                    f"Translating tables... ({processed}/{total_elements})"
                )

        # 3. Translate headers and footers
        for section in doc.sections:
            # Header
            for para in section.header.paragraphs:
                if is_translatable(para.text):
                    translated = translator.translate_text(para.text)
                    self._set_paragraph_text(para, translated)

            # First page header
            if section.first_page_header:
                for para in section.first_page_header.paragraphs:
                    if is_translatable(para.text):
                        translated = translator.translate_text(para.text)
                        self._set_paragraph_text(para, translated)

            # Footer
            for para in section.footer.paragraphs:
                if is_translatable(para.text):
                    translated = translator.translate_text(para.text)
                    self._set_paragraph_text(para, translated)

            # First page footer
            if section.first_page_footer:
                for para in section.first_page_footer.paragraphs:
                    if is_translatable(para.text):
                        translated = translator.translate_text(para.text)
                        self._set_paragraph_text(para, translated)

            processed += 1

        # 4. Translate text boxes (these are in the XML)
        self._translate_text_boxes(doc, translator)

        # Save
        doc.save(output_path)
        if progress_callback:
            progress_callback(1.0, "DOCX translation complete!")

    def _set_paragraph_text(self, para, translated_text):
        """
        Set the text of a paragraph while preserving the first run's formatting.
        Clears all runs except the first, puts translated text in first run.
        """
        if not para.runs:
            # No runs exist — just set the text directly
            para.text = translated_text
            return

        # Keep first run's formatting, set translated text
        first_run = para.runs[0]
        first_run.text = translated_text

        # Clear remaining runs
        for run in para.runs[1:]:
            run.text = ""

    def _translate_text_boxes(self, doc, translator):
        """
        Translate text inside text boxes (which are stored as shapes in XML).
        This is a best-effort approach.
        """
        try:
            from docx.oxml.ns import qn

            for textbox in doc.element.findall('.//' + qn('w:txbxContent')):
                for para_elem in textbox.findall(qn('w:p')):
                    for run_elem in para_elem.findall(qn('w:r')):
                        text_elem = run_elem.find(qn('w:t'))
                        if text_elem is not None and text_elem.text and is_translatable(text_elem.text):
                            text_elem.text = translator.translate_text(text_elem.text)
        except Exception:
            pass  # Text boxes are best-effort

    def _count_elements(self, doc):
        """Count total translatable elements for progress tracking."""
        count = len(doc.paragraphs)
        for table in doc.tables:
            count += len(table.rows) * len(table.columns)
        count += len(doc.sections) * 4  # headers/footers
        return max(count, 1)